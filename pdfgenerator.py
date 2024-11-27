from flask import Flask, render_template, request, jsonify, send_from_directory
from docx import Document
import os
import pythoncom
import comtypes.client

app = Flask(__name__)

# Set up paths
base_dir = os.path.abspath(os.path.dirname(__file__))
STATIC_DIR = os.path.join(base_dir, "static")
os.makedirs(STATIC_DIR, exist_ok=True)

# Function to edit the Word template
def edit_word_template(template_path, output_path, name, designation, contact, email, location, selected_services):
    try:
        doc = Document(template_path)
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            if "<<Client Name>>" in para.text:
                para.text = para.text.replace("<<Client Name>>", name)
            if "<<Client Designation>>" in para.text:
                para.text = para.text.replace("<<Client Designation>>", designation)
            if "<<Client Contact>>" in para.text:
                para.text = para.text.replace("<<Client Contact>>", contact)
            if "<<Client Email>>" in para.text:
                para.text = para.text.replace("<<Client Email>>", email)
            if "<<Client Location>>" in para.text:
                para.text = para.text.replace("<<Client Location>>", location)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "<<Client Name>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Name>>", name)
                    if "<<Client Designation>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Designation>>", designation)
                    if "<<Client Contact>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Contact>>", contact)
                    if "<<Client Email>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Email>>", email)
                    if "<<Client Location>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Location>>", location)

        # Filter rows based on selected services
        for table in doc.tables:
            for row in table.rows[1:]:  # Skip the header row
                service_name = row.cells[0].text.strip()
                if service_name not in selected_services:
                    row._element.getparent().remove(row._element)

        # Save the updated document
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")


# Function to convert Word to PDF
def convert_to_pdf(doc_path, pdf_path):
    word = None
    try:
        pythoncom.CoInitialize()
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
    except Exception as e:
        raise Exception(f"Error converting Word to PDF: {e}")
    finally:
        if word:
            word.Quit()
        pythoncom.CoUninitialize()


@app.route("/")
def home():
    return render_template("home.html")


@app.route("/generate-pdf", methods=["POST"])
def generate_pdf():
    try:
        # File paths
        template_path = os.path.join(base_dir, "DM & Automations Services Pricing - Andrew.docx")
        word_output_path = os.path.join(STATIC_DIR, "Customized_Pricing.docx")
        pdf_output_path = os.path.join(STATIC_DIR, "Customized_Pricing.pdf")

        # Gather form data
        name = request.form["name"]
        designation = request.form["designation"]
        contact = request.form["contact"]
        email = request.form["email"]
        location = request.form["location"]
        selected_services = request.form.getlist("selected_services")

        # Validate input
        if not all([name, designation, contact, email, location, selected_services]):
            return jsonify({"success": False, "error": "All fields are required and at least one service must be selected."})

        # Process Word template and generate PDF
        edit_word_template(template_path, word_output_path, name, designation, contact, email, location, selected_services)
        convert_to_pdf(word_output_path, pdf_output_path)

        return jsonify({"success": True, "pdf_url": f"/static/Customized_Pricing.pdf"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/static/<path:filename>")
def serve_static(filename):
    return send_from_directory(STATIC_DIR, filename)


if __name__ == "__main__":
    app.run(debug=True)
